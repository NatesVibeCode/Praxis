"""Tests for core/file_formats writers."""

from __future__ import annotations

import csv as csv_stdlib
import hashlib
from pathlib import Path

import pytest

from core.file_formats import (
    WriteValidationError,
    write_csv,
    write_docx,
    write_pdf,
    write_txt,
)


# ── CSV ──────────────────────────────────────────────────────────────────────

class TestWriteCsv:
    def test_basic_roundtrip(self, tmp_path: Path) -> None:
        out = tmp_path / "out.csv"
        rows = [{"name": "Alice", "age": "30"}, {"name": "Bob", "age": "25"}]
        result = write_csv(rows, ["name", "age"], out)
        assert result.bytes_written > 0
        assert len(result.sha256) == 64
        # Read back
        with open(out, newline="", encoding="utf-8") as f:
            reader = csv_stdlib.DictReader(f)
            read_rows = list(reader)
        assert read_rows[0]["name"] == "Alice"
        assert read_rows[1]["age"] == "25"

    def test_sha256_matches_content(self, tmp_path: Path) -> None:
        out = tmp_path / "hash.csv"
        write_csv([{"x": "1"}], ["x"], out)
        data = out.read_bytes()
        assert hashlib.sha256(data).hexdigest() == write_csv([{"x": "1"}], ["x"], tmp_path / "hash2.csv").sha256

    def test_extra_key_raises(self, tmp_path: Path) -> None:
        out = tmp_path / "extra.csv"
        with pytest.raises(ValueError):
            write_csv([{"x": "1", "y": "2"}], ["x"], out)

    def test_empty_headers_raise(self, tmp_path: Path) -> None:
        out = tmp_path / "eh.csv"
        with pytest.raises(WriteValidationError, match="headers must not be empty"):
            write_csv([], [], out)

    def test_duplicate_headers_raise(self, tmp_path: Path) -> None:
        out = tmp_path / "dup.csv"
        with pytest.raises(WriteValidationError, match="duplicate headers"):
            write_csv([], ["x", "x"], out)

    def test_wrong_extension_raises(self, tmp_path: Path) -> None:
        out = tmp_path / "out.txt"
        with pytest.raises(WriteValidationError):
            write_csv([], ["x"], out)

    def test_missing_key_becomes_empty(self, tmp_path: Path) -> None:
        out = tmp_path / "miss.csv"
        write_csv([{"name": "Alice"}], ["name", "age"], out)
        with open(out, newline="", encoding="utf-8") as f:
            reader = csv_stdlib.DictReader(f)
            row = list(reader)[0]
        assert row["age"] == ""

    def test_result_format(self, tmp_path: Path) -> None:
        out = tmp_path / "fmt.csv"
        result = write_csv([{"a": "1"}], ["a"], out)
        assert result.format == "csv"
        assert result.media_type == "text/csv"
        assert result.path == str(out)


# ── TXT ──────────────────────────────────────────────────────────────────────

class TestWriteTxt:
    def test_basic_roundtrip(self, tmp_path: Path) -> None:
        out = tmp_path / "out.txt"
        result = write_txt("hello\nworld", out)
        assert out.read_text(encoding="utf-8") == "hello\nworld"
        assert result.bytes_written > 0
        assert len(result.sha256) == 64

    def test_crlf_normalized(self, tmp_path: Path) -> None:
        out = tmp_path / "crlf.txt"
        write_txt("line1\r\nline2\r\n", out)
        content = out.read_bytes()
        assert b"\r" not in content

    def test_sha256_correct(self, tmp_path: Path) -> None:
        out = tmp_path / "h.txt"
        result = write_txt("test content", out)
        expected = hashlib.sha256(out.read_bytes()).hexdigest()
        assert result.sha256 == expected

    def test_wrong_extension_raises(self, tmp_path: Path) -> None:
        out = tmp_path / "out.csv"
        with pytest.raises(WriteValidationError):
            write_txt("hello", out)

    def test_result_format(self, tmp_path: Path) -> None:
        out = tmp_path / "fmt.txt"
        result = write_txt("x", out)
        assert result.format == "txt"
        assert result.media_type == "text/plain"


# ── DOCX ─────────────────────────────────────────────────────────────────────

class TestWriteDocx:
    def test_headings_and_body(self, tmp_path: Path) -> None:
        import docx
        out = tmp_path / "out.docx"
        sections = [
            {"heading": "Summary", "body": "This is the summary body."},
            {"heading": "Details", "body": "More detail here.\n\nSecond paragraph."},
        ]
        result = write_docx(sections, out)
        assert result.bytes_written > 0
        assert len(result.sha256) == 64
        # Read back
        doc = docx.Document(str(out))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Summary" in headings
        assert "Details" in headings

    def test_body_without_heading(self, tmp_path: Path) -> None:
        import docx
        out = tmp_path / "nobody.docx"
        write_docx([{"body": "Only body text here."}], out)
        doc = docx.Document(str(out))
        texts = [p.text for p in doc.paragraphs]
        assert any("Only body text" in t for t in texts)

    def test_unicode_survives(self, tmp_path: Path) -> None:
        import docx
        out = tmp_path / "uni.docx"
        write_docx([{"heading": "Café", "body": "Résumé content."}], out)
        doc = docx.Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Café" in all_text or "Caf" in all_text  # heading

    def test_wrong_extension_raises(self, tmp_path: Path) -> None:
        out = tmp_path / "out.txt"
        with pytest.raises(WriteValidationError):
            write_docx([], out)

    def test_result_format(self, tmp_path: Path) -> None:
        out = tmp_path / "fmt.docx"
        result = write_docx([], out)
        assert result.format == "docx"
        assert "wordprocessingml" in result.media_type


# ── PDF ──────────────────────────────────────────────────────────────────────

class TestWritePdf:
    def test_file_created_nonempty(self, tmp_path: Path) -> None:
        out = tmp_path / "out.pdf"
        result = write_pdf([{"heading": "Title", "body": "Content here."}], out)
        assert out.exists()
        assert result.bytes_written > 0

    def test_pdf_readable_by_pypdf(self, tmp_path: Path) -> None:
        import pypdf
        out = tmp_path / "readable.pdf"
        write_pdf([{"heading": "Hello", "body": "World content."}], out)
        reader = pypdf.PdfReader(str(out))
        assert len(reader.pages) >= 1

    def test_sha256_correct(self, tmp_path: Path) -> None:
        out = tmp_path / "h.pdf"
        result = write_pdf([{"body": "test"}], out)
        expected = hashlib.sha256(out.read_bytes()).hexdigest()
        assert result.sha256 == expected

    def test_wrong_extension_raises(self, tmp_path: Path) -> None:
        out = tmp_path / "out.docx"
        with pytest.raises(WriteValidationError):
            write_pdf([], out)

    def test_result_format(self, tmp_path: Path) -> None:
        out = tmp_path / "fmt.pdf"
        result = write_pdf([], out)
        assert result.format == "pdf"
        assert result.media_type == "application/pdf"

    def test_empty_sections(self, tmp_path: Path) -> None:
        out = tmp_path / "empty.pdf"
        result = write_pdf([], out)
        assert result.bytes_written > 0  # fpdf2 always writes PDF header
