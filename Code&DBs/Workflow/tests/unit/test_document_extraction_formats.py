"""Backward-compatibility tests for memory/document_extraction.py."""

from __future__ import annotations

from pathlib import Path

import pytest

from memory.document_extraction import DocumentExtractor, ExtractedDocument


class TestDocumentExtractorBackwardCompat:
    def test_supported_formats_include_pdf_docx(self) -> None:
        ex = DocumentExtractor()
        assert ".pdf" in ex.supported_formats()
        assert ".docx" in ex.supported_formats()

    def test_csv_returns_extracted_document(self, tmp_path: Path) -> None:
        f = tmp_path / "data.csv"
        f.write_text('name,score\nAlice,90\nBob,80\n', encoding="utf-8")
        doc = DocumentExtractor().extract(str(f))
        assert isinstance(doc, ExtractedDocument)
        assert doc.format == ".csv"
        assert "name" in doc.sections
        assert "score" in doc.sections
        assert doc.content  # non-empty
        assert "column_count" in doc.metadata

    def test_csv_quoted_fields_handled(self, tmp_path: Path) -> None:
        f = tmp_path / "q.csv"
        f.write_text('id,val\n1,"hello, world"\n', encoding="utf-8")
        doc = DocumentExtractor().extract(str(f))
        # Should not crash; structured data should be present
        assert isinstance(doc, ExtractedDocument)

    def test_txt_backward_compat(self, tmp_path: Path) -> None:
        f = tmp_path / "note.txt"
        f.write_text("some text content", encoding="utf-8")
        doc = DocumentExtractor().extract(str(f))
        assert doc.format == ".txt"
        assert doc.content == "some text content"
        assert "char_count" in doc.metadata

    def test_pdf_dispatches_correctly(self, tmp_path: Path) -> None:
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.cell(0, 10, "Test PDF content")
        out = tmp_path / "doc.pdf"
        out.write_bytes(bytes(pdf.output()))
        doc = DocumentExtractor().extract(str(out))
        assert isinstance(doc, ExtractedDocument)
        assert doc.format == ".pdf"
        assert "page_count" in doc.metadata
        assert "warnings" in doc.metadata
        assert "truncated" in doc.metadata

    def test_docx_dispatches_correctly(self, tmp_path: Path) -> None:
        import docx
        d = docx.Document()
        d.add_heading("Test Heading", level=1)
        d.add_paragraph("Some paragraph text.")
        out = tmp_path / "doc.docx"
        d.save(str(out))
        doc = DocumentExtractor().extract(str(out))
        assert isinstance(doc, ExtractedDocument)
        assert doc.format == ".docx"
        assert "Test Heading" in doc.sections
        assert "paragraph_count" in doc.metadata

    def test_unsupported_extension_raises(self, tmp_path: Path) -> None:
        f = tmp_path / "file.xyz"
        f.write_bytes(b"fake content")
        with pytest.raises(ValueError, match="Unsupported format"):
            DocumentExtractor().extract(str(f))

    def test_extract_batch_skips_errors(self, tmp_path: Path) -> None:
        good = tmp_path / "good.txt"
        good.write_text("ok", encoding="utf-8")
        missing = str(tmp_path / "missing.txt")
        results = DocumentExtractor().extract_batch([str(good), missing])
        assert len(results) == 1
        assert results[0].format == ".txt"
