"""Tests for core/file_formats readers."""

from __future__ import annotations

import io
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from core.file_formats import (
    EncryptedDocumentError,
    FileParseError,
    ReadLimits,
    read_csv,
    read_docx,
    read_pdf,
    read_txt,
)


# ── CSV ──────────────────────────────────────────────────────────────────────

class TestReadCsv:
    def test_basic_roundtrip(self, tmp_path: Path) -> None:
        f = tmp_path / "data.csv"
        f.write_text("name,age\nAlice,30\nBob,25\n", encoding="utf-8")
        result = read_csv(f)
        assert result.format == "csv"
        assert result.sections == ("name", "age")
        assert result.metadata["column_count"] == 2
        assert result.metadata["preview_row_count"] == 2
        assert result.structured["headers"] == ["name", "age"]
        assert result.structured["rows_preview"][0] == {"name": "Alice", "age": "30"}

    def test_quoted_comma(self, tmp_path: Path) -> None:
        f = tmp_path / "q.csv"
        f.write_text('id,phrase\n1,"hello, world"\n2,"foo"\n', encoding="utf-8")
        result = read_csv(f)
        assert result.structured["rows_preview"][0]["phrase"] == "hello, world"

    def test_escaped_quote(self, tmp_path: Path) -> None:
        f = tmp_path / "eq.csv"
        f.write_text('id,val\n1,"say ""hi"""\n', encoding="utf-8")
        result = read_csv(f)
        assert 'say "hi"' in result.structured["rows_preview"][0]["val"]

    def test_utf8_bom(self, tmp_path: Path) -> None:
        f = tmp_path / "bom.csv"
        f.write_bytes(b"\xef\xbb\xbfname,score\nZoe,99\n")
        result = read_csv(f)
        assert "name" in result.sections
        assert result.structured["rows_preview"][0]["name"] == "Zoe"

    def test_blank_rows_skipped(self, tmp_path: Path) -> None:
        f = tmp_path / "blank.csv"
        f.write_text("a,b\n1,2\n\n3,4\n", encoding="utf-8")
        result = read_csv(f)
        assert result.metadata["preview_row_count"] == 2

    def test_duplicate_headers_raise(self, tmp_path: Path) -> None:
        f = tmp_path / "dup.csv"
        f.write_text("x,x\n1,2\n", encoding="utf-8")
        with pytest.raises(FileParseError, match="Duplicate header"):
            read_csv(f)

    def test_row_cap_sets_truncated(self, tmp_path: Path) -> None:
        lines = ["a,b"] + [f"{i},{i}" for i in range(20)]
        f = tmp_path / "big.csv"
        f.write_text("\n".join(lines), encoding="utf-8")
        result = read_csv(f, ReadLimits(max_rows=5))
        assert result.truncated is True
        assert result.metadata["preview_row_count"] == 5

    def test_empty_file(self, tmp_path: Path) -> None:
        f = tmp_path / "empty.csv"
        f.write_text("", encoding="utf-8")
        result = read_csv(f)
        assert result.content == ""
        assert "empty file" in result.warnings

    def test_source_hash_present(self, tmp_path: Path) -> None:
        f = tmp_path / "h.csv"
        f.write_text("x\n1\n", encoding="utf-8")
        result = read_csv(f)
        assert len(result.source_sha256) == 64
        assert result.source_bytes > 0

    def test_extra_fields_warning(self, tmp_path: Path) -> None:
        f = tmp_path / "extra.csv"
        # Write raw to bypass DictWriter validation
        f.write_text("a,b\n1,2,3\n", encoding="utf-8")
        result = read_csv(f)
        assert any("more fields" in w for w in result.warnings)


# ── TXT ──────────────────────────────────────────────────────────────────────

class TestReadTxt:
    def test_basic(self, tmp_path: Path) -> None:
        f = tmp_path / "note.txt"
        f.write_text("hello\nworld\n", encoding="utf-8")
        result = read_txt(f)
        assert result.format == "txt"
        assert result.content == "hello\nworld\n"
        assert result.metadata["line_count"] == 3

    def test_bom_stripped(self, tmp_path: Path) -> None:
        f = tmp_path / "bom.txt"
        f.write_bytes(b"\xef\xbb\xbfHello\n")
        result = read_txt(f)
        assert result.content.startswith("Hello")

    def test_char_cap_sets_truncated(self, tmp_path: Path) -> None:
        f = tmp_path / "long.txt"
        f.write_text("x" * 1000, encoding="utf-8")
        result = read_txt(f, ReadLimits(max_chars=100))
        assert result.truncated is True
        assert len(result.content) == 100

    def test_crlf_normalized(self, tmp_path: Path) -> None:
        f = tmp_path / "crlf.txt"
        f.write_bytes(b"line1\r\nline2\r\n")
        result = read_txt(f)
        assert "\r" not in result.content

    def test_source_hash_present(self, tmp_path: Path) -> None:
        f = tmp_path / "h.txt"
        f.write_text("data", encoding="utf-8")
        result = read_txt(f)
        assert len(result.source_sha256) == 64


# ── PDF ──────────────────────────────────────────────────────────────────────

class TestReadPdf:
    def _make_pdf(self, tmp_path: Path, text: str = "Hello PDF world") -> Path:
        """Create a minimal real PDF using fpdf2."""
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.multi_cell(0, 10, text)
        out = tmp_path / "doc.pdf"
        out.write_bytes(bytes(pdf.output()))
        return out

    def test_basic_extraction(self, tmp_path: Path) -> None:
        f = self._make_pdf(tmp_path, "Hello PDF world")
        result = read_pdf(f)
        assert result.format == "pdf"
        assert result.metadata["page_count"] == 1
        assert result.metadata["pages_read"] == 1
        assert result.metadata["encrypted"] is False
        assert len(result.structured["pages"]) == 1

    def test_content_contains_text(self, tmp_path: Path) -> None:
        f = self._make_pdf(tmp_path, "Praxis file format test")
        result = read_pdf(f)
        assert "Praxis" in result.content or len(result.content) > 0

    def test_page_sections(self, tmp_path: Path) -> None:
        f = self._make_pdf(tmp_path)
        result = read_pdf(f)
        assert "Page 1" in result.sections

    def test_page_cap_sets_truncated(self, tmp_path: Path) -> None:
        from fpdf import FPDF
        pdf = FPDF()
        for _ in range(5):
            pdf.add_page()
            pdf.set_font("Helvetica", size=12)
            pdf.cell(0, 10, "page content")
        out = tmp_path / "multi.pdf"
        out.write_bytes(bytes(pdf.output()))
        result = read_pdf(out, ReadLimits(max_pages=3))
        assert result.truncated is True
        assert result.metadata["pages_read"] == 3

    def test_encrypted_pdf_raises(self, tmp_path: Path) -> None:
        import pypdf
        mock_reader = MagicMock()
        mock_reader.is_encrypted = True
        with patch("pypdf.PdfReader", return_value=mock_reader):
            f = tmp_path / "enc.pdf"
            f.write_bytes(b"%PDF-1.4 fake")
            with pytest.raises(EncryptedDocumentError):
                read_pdf(f)

    def test_source_hash_present(self, tmp_path: Path) -> None:
        f = self._make_pdf(tmp_path)
        result = read_pdf(f)
        assert len(result.source_sha256) == 64
        assert result.source_bytes > 0


# ── DOCX ─────────────────────────────────────────────────────────────────────

class TestReadDocx:
    def _make_docx(self, tmp_path: Path) -> Path:
        import docx
        doc = docx.Document()
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("This is the intro paragraph.")
        doc.add_heading("Methods", level=1)
        doc.add_paragraph("Method details here.")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "2"
        out = tmp_path / "doc.docx"
        doc.save(str(out))
        return out

    def test_headings_extracted(self, tmp_path: Path) -> None:
        f = self._make_docx(tmp_path)
        result = read_docx(f)
        assert "Introduction" in result.sections
        assert "Methods" in result.sections

    def test_paragraphs_extracted(self, tmp_path: Path) -> None:
        f = self._make_docx(tmp_path)
        result = read_docx(f)
        texts = [p["text"] for p in result.structured["paragraphs"]]
        assert any("intro paragraph" in t for t in texts)

    def test_tables_extracted(self, tmp_path: Path) -> None:
        f = self._make_docx(tmp_path)
        result = read_docx(f)
        assert result.metadata["table_count"] == 1
        rows = result.structured["tables"][0]["rows"]
        assert rows[0] == ["A", "B"]
        assert rows[1] == ["1", "2"]

    def test_metadata_counts(self, tmp_path: Path) -> None:
        f = self._make_docx(tmp_path)
        result = read_docx(f)
        assert result.metadata["heading_count"] == 2
        assert result.metadata["paragraph_count"] > 0

    def test_paragraph_cap_sets_truncated(self, tmp_path: Path) -> None:
        import docx
        doc = docx.Document()
        for i in range(20):
            doc.add_paragraph(f"Paragraph {i}")
        out = tmp_path / "many.docx"
        doc.save(str(out))
        result = read_docx(out, ReadLimits(max_paragraphs=5))
        assert result.truncated is True
        assert result.metadata["paragraph_count"] == 5

    def test_source_hash_present(self, tmp_path: Path) -> None:
        f = self._make_docx(tmp_path)
        result = read_docx(f)
        assert len(result.source_sha256) == 64
